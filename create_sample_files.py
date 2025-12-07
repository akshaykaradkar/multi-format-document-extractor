"""
Script to generate sample files for all 5 clients.
Run this once to create the mock documents.
"""

import os
from pathlib import Path

SAMPLE_DIR = Path(__file__).parent / "sample_data"
SAMPLE_DIR.mkdir(exist_ok=True)


def create_pdf_sample():
    """Create Client A - TechCorp Solutions PDF."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        pdf_path = SAMPLE_DIR / "client_a_techcorp.pdf"
        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        # Header
        elements.append(Paragraph("TechCorp Solutions", styles['Title']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("<b>PURCHASE ORDER #PO-2024-1247</b>", styles['Heading2']))
        elements.append(Spacer(1, 12))

        # Order Info
        info_data = [
            ["Date:", "March 15, 2024"],
            ["Delivery Required:", "March 22, 2024"],
        ]
        info_table = Table(info_data, colWidths=[150, 200])
        elements.append(info_table)
        elements.append(Spacer(1, 20))

        # Items Table
        items_data = [
            ["Item Code", "Description", "Qty", "Unit Price", "Total"],
            ["TC-001", "Widget Pro", "50", "$25.00", "$1,250.00"],
            ["TC-002", "Gadget Max", "25", "$45.00", "$1,125.00"],
            ["", "", "", "TOTAL:", "$2,375.00"],
        ]
        items_table = Table(items_data, colWidths=[80, 150, 50, 80, 90])
        items_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(items_table)
        elements.append(Spacer(1, 20))

        # Special Notes
        elements.append(Paragraph("<b>Special Notes:</b> Rush delivery required", styles['Normal']))

        doc.build(elements)
        print(f"Created: {pdf_path}")
        return True

    except ImportError:
        print("reportlab not installed. Creating text-based PDF alternative...")
        # Create a simple text file as fallback
        txt_path = SAMPLE_DIR / "client_a_techcorp.txt"
        content = """TechCorp Solutions
PURCHASE ORDER #PO-2024-1247

Date: March 15, 2024
Delivery Required: March 22, 2024

Item Code | Description | Qty | Unit Price | Total
TC-001 | Widget Pro | 50 | $25.00 | $1,250.00
TC-002 | Gadget Max | 25 | $45.00 | $1,125.00

TOTAL: $2,375.00

Special Notes: Rush delivery required
"""
        txt_path.write_text(content)
        print(f"Created text fallback: {txt_path}")
        return False


def create_excel_sample():
    """Create Client B - Global Manufacturing Inc Excel."""
    try:
        import pandas as pd

        xlsx_path = SAMPLE_DIR / "client_b_global_mfg.xlsx"

        # Create multiple sheets
        with pd.ExcelWriter(xlsx_path, engine='openpyxl') as writer:
            # Sheet 1: Order_Info
            order_info = pd.DataFrame({
                'Order#': ['GMI-2024-0892'],
                'Client_Name': ['Global Manufacturing Inc'],
                'Order_Created': ['2024-03-18'],
                'Needed_By': ['2024-03-25']
            })
            order_info.to_excel(writer, sheet_name='Order_Info', index=False)

            # Sheet 2: Line_Items
            line_items = pd.DataFrame({
                'SKU': ['GMI-SKU-001', 'GMI-SKU-002', 'GMI-SKU-003'],
                'Item_Desc': ['Precision Bearing Assembly', 'Hydraulic Pump Unit', 'Control Valve Set'],
                'Order_Qty': [200, 50, 100],
                'Price_Each': [15.75, 89.00, 34.50]
            })
            line_items.to_excel(writer, sheet_name='Line_Items', index=False)

            # Sheet 3: Notes
            notes = pd.DataFrame({
                'Special_Requirements': ['Temperature controlled packaging required'],
                'Delivery_Instructions': ['Dock B, Building 3. Call 30 min before arrival.']
            })
            notes.to_excel(writer, sheet_name='Notes', index=False)

        print(f"Created: {xlsx_path}")
        return True

    except ImportError as e:
        print(f"pandas/openpyxl not installed: {e}")
        return False


def create_word_sample():
    """Create Client C - Regional Distributors Word doc."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        docx_path = SAMPLE_DIR / "client_c_regional_dist.docx"
        doc = Document()

        # Title
        title = doc.add_heading('Order Request - Regional Distributors', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Order Number
        doc.add_paragraph('Order Number: RD-240815-A')
        doc.add_paragraph('')

        # Free text section
        doc.add_paragraph('We need the following items by August 22, 2024:')
        doc.add_paragraph('')

        # Item 1
        doc.add_paragraph('Product Name: Industrial Pump Model X200')
        doc.add_paragraph('Quantity Needed: 3 units')
        doc.add_paragraph('Expected Price: $850 per unit')
        doc.add_paragraph('')

        # Item 2
        doc.add_paragraph('Product Name: Filter Cartridge Set')
        doc.add_paragraph('Quantity Needed: 12 sets')
        doc.add_paragraph('Expected Price: $45 per set')
        doc.add_paragraph('')

        # Notes
        doc.add_paragraph('Please note: This is for our Phoenix warehouse.')
        doc.add_paragraph('Delivery must be completed before 3 PM.')

        doc.save(str(docx_path))
        print(f"Created: {docx_path}")
        return True

    except ImportError as e:
        print(f"python-docx not installed: {e}")
        return False


def main():
    """Create all sample files."""
    print("Creating sample files for Document Automation PoC...")
    print("=" * 50)

    results = {
        "PDF (Client A)": create_pdf_sample(),
        "Excel (Client B)": create_excel_sample(),
        "Word (Client C)": create_word_sample(),
    }

    # CSV is already created
    csv_path = SAMPLE_DIR / "client_d_supply_chain.csv"
    results["CSV (Client D)"] = csv_path.exists()

    # Scanned image
    img_path = SAMPLE_DIR / "client_e_local_hardware.jpg"
    results["Image (Client E)"] = img_path.exists()

    print("\n" + "=" * 50)
    print("Summary:")
    for name, success in results.items():
        status = "[OK]" if success else "[FAIL]"
        print(f"  {status} {name}")

    print("\nSample files are in:", SAMPLE_DIR)


if __name__ == "__main__":
    main()
