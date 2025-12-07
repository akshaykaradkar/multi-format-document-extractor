"""
Word Document Parser for mixed-format documents.

Handles Client C (Regional Distributors) - Word document with
embedded tables, free-text sections, and inconsistent structure.
"""

import re
from pathlib import Path
from typing import Union

from docx import Document

from .base_parser import BaseParser, ParseError
from ..schemas import RawExtraction


class WordParser(BaseParser):
    """
    Parser for Word documents with mixed formatting.

    Handles embedded tables and free-text order information
    with inconsistent formatting.
    """

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self) -> RawExtraction:
        """
        Parse Word document and extract order data.

        Returns:
            RawExtraction with parsed order data
        """
        try:
            doc = Document(self.file_path)

            # Extract all text from paragraphs
            all_text = "\n".join(p.text for p in doc.paragraphs)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

            return self._extract_order_data(all_text, tables)

        except Exception as e:
            raise ParseError(
                f"Failed to parse Word document: {str(e)}",
                file_path=self.file_path,
                original_error=e,
            )

    def _extract_order_data(
        self, text: str, tables: list[list[list[str]]]
    ) -> RawExtraction:
        """
        Extract order data from Word document content.

        Args:
            text: All paragraph text from document
            tables: List of tables (each table is list of rows)

        Returns:
            RawExtraction object
        """
        # Extract order ID
        order_id = self._extract_pattern(
            text,
            [
                r"Order\s*(?:Number|No\.?|#)?\s*:?\s*([A-Z0-9\-]+)",
                r"PO\s*(?:Number|No\.?|#)?\s*:?\s*([A-Z0-9\-]+)",
                r"Reference\s*:?\s*([A-Z0-9\-]+)",
            ],
        )

        # Extract client name from header
        client_name = self._extract_client_name(text)

        # Extract dates
        delivery_date = self._extract_date(
            text,
            [
                r"(?:by|before|need\s*(?:by)?|delivery)\s*([A-Za-z]+\s+\d{1,2},?\s+\d{4})",
                r"(?:by|before|need\s*(?:by)?|delivery)\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})",
            ],
        )

        order_date = self._extract_date(
            text,
            [
                r"(?:Date|Dated|Created)\s*:?\s*([A-Za-z]+\s+\d{1,2},?\s+\d{4})",
                r"(?:Date|Dated|Created)\s*:?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})",
            ],
        )

        # If no order date found, use current date
        if not order_date:
            from datetime import datetime
            order_date = datetime.now().strftime("%Y-%m-%d")

        # Extract items - try tables first, then free text
        items = self._extract_items_from_tables(tables)
        if not items:
            items = self._extract_items_from_text(text)

        # Calculate total
        order_total = sum(item.get("total_price", 0) for item in items)

        # Extract special instructions
        special_instructions = self._extract_special_instructions(text)

        return RawExtraction(
            order_id=order_id,
            client_name=client_name,
            order_date=order_date,
            delivery_date=delivery_date,
            items=items,
            order_total=order_total,
            currency="USD",
            special_instructions=special_instructions,
            source_confidence=0.85,  # Medium-high confidence for Word docs
            extraction_method="word_python_docx",
        )

    def _extract_pattern(self, text: str, patterns: list[str]) -> str | None:
        """Try multiple regex patterns and return first match."""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                return match.group(1).strip()
        return None

    def _extract_client_name(self, text: str) -> str | None:
        """Extract client/company name from document."""
        # Look in header/title
        lines = text.split("\n")

        for line in lines[:10]:
            line = line.strip()
            # Pattern for "Order Request - Company Name"
            match = re.match(r"Order\s*Request\s*[-–]\s*(.+)", line, re.IGNORECASE)
            if match:
                return match.group(1).strip()

            # Pattern for standalone company names
            if any(
                suffix in line
                for suffix in [
                    "Distributors",
                    "Inc",
                    "Corp",
                    "LLC",
                    "Ltd",
                    "Co.",
                    "Company",
                    "Partners",
                ]
            ):
                return line

        return None

    def _extract_date(self, text: str, patterns: list[str]) -> str | None:
        """Extract and normalize date from text."""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return self._normalize_date(match.group(1))
        return None

    def _normalize_date(self, date_str: str) -> str:
        """Convert various date formats to YYYY-MM-DD."""
        from dateutil import parser

        try:
            parsed = parser.parse(date_str)
            return parsed.strftime("%Y-%m-%d")
        except (ValueError, TypeError):
            return date_str

    def _extract_items_from_tables(self, tables: list[list[list[str]]]) -> list[dict]:
        """Extract items from Word tables."""
        items = []

        for table in tables:
            if len(table) < 2:
                continue

            # Check if this looks like an items table
            header = table[0]
            header_text = " ".join(str(cell).lower() for cell in header)

            if not any(
                kw in header_text
                for kw in ["product", "item", "qty", "quantity", "price", "description"]
            ):
                continue

            # Map columns
            col_map = self._map_columns(header)

            # Parse data rows
            for row in table[1:]:
                item = self._parse_table_row(row, col_map)
                if item and item.get("description"):
                    items.append(item)

        return items

    def _map_columns(self, header: list[str]) -> dict[str, int]:
        """Map column indices to field names."""
        col_map = {}
        keywords = {
            "product_code": ["code", "sku", "product code"],
            "description": ["description", "product", "item", "name"],
            "quantity": ["qty", "quantity", "amount"],
            "unit_price": ["price", "unit price", "rate", "cost"],
            "total_price": ["total", "amount", "line total"],
        }

        for i, cell in enumerate(header):
            cell_lower = str(cell).lower().strip()
            for field, kws in keywords.items():
                if any(kw in cell_lower for kw in kws):
                    col_map[field] = i
                    break

        return col_map

    def _parse_table_row(self, row: list[str], col_map: dict[str, int]) -> dict | None:
        """Parse a table row into an item dict."""
        try:
            item = {}

            # Get description
            if "description" in col_map and col_map["description"] < len(row):
                desc = str(row[col_map["description"]]).strip()
                if desc:
                    item["description"] = desc
                else:
                    return None
            else:
                return None

            # Get product code
            if "product_code" in col_map and col_map["product_code"] < len(row):
                item["product_code"] = str(row[col_map["product_code"]]).strip()
            else:
                item["product_code"] = self._generate_code(item["description"])

            # Get quantity
            if "quantity" in col_map and col_map["quantity"] < len(row):
                qty_str = str(row[col_map["quantity"]])
                item["quantity"] = self._parse_number(qty_str, as_int=True)
            else:
                item["quantity"] = 1

            # Get unit price
            if "unit_price" in col_map and col_map["unit_price"] < len(row):
                price_str = str(row[col_map["unit_price"]])
                item["unit_price"] = self._parse_number(price_str)
            else:
                item["unit_price"] = 0.0

            # Get or calculate total
            if "total_price" in col_map and col_map["total_price"] < len(row):
                item["total_price"] = self._parse_number(str(row[col_map["total_price"]]))
            else:
                item["total_price"] = item["quantity"] * item["unit_price"]

            return item

        except (ValueError, IndexError):
            return None

    def _extract_items_from_text(self, text: str) -> list[dict]:
        """Extract items from free-text format."""
        items = []

        # Pattern for free-text item format:
        # "Product Name: Industrial Pump Model X200
        #  Quantity Needed: 3 units
        #  Expected Price: $850 per unit"
        blocks = re.split(
            r"(?=Product\s*Name|Item\s*:)", text, flags=re.IGNORECASE
        )

        for block in blocks:
            if not block.strip():
                continue

            item = {}

            # Extract product name/description
            name_match = re.search(
                r"(?:Product\s*Name|Item)\s*:?\s*(.+?)(?:\n|Quantity|Expected|$)",
                block,
                re.IGNORECASE | re.DOTALL,
            )
            if name_match:
                item["description"] = name_match.group(1).strip()

            # Extract quantity
            qty_match = re.search(
                r"Quantity\s*(?:Needed|Required)?\s*:?\s*(\d+)\s*(?:units?|pcs?|sets?)?",
                block,
                re.IGNORECASE,
            )
            if qty_match:
                item["quantity"] = int(qty_match.group(1))
            else:
                item["quantity"] = 1

            # Extract price
            price_match = re.search(
                r"(?:Expected\s*)?Price\s*:?\s*\$?([\d,\.]+)\s*(?:per|each|/)?",
                block,
                re.IGNORECASE,
            )
            if price_match:
                item["unit_price"] = float(price_match.group(1).replace(",", ""))
            else:
                item["unit_price"] = 0.0

            if item.get("description"):
                item["product_code"] = self._generate_code(item["description"])
                item["total_price"] = item["quantity"] * item["unit_price"]
                items.append(item)

        return items

    def _extract_special_instructions(self, text: str) -> str | None:
        """Extract special instructions from text."""
        instructions = []

        # Look for "Please note:", "Note:", etc.
        note_match = re.search(
            r"(?:Please\s*)?[Nn]ote[s]?\s*:?\s*(.+?)(?:\n\n|\Z)",
            text,
            re.DOTALL,
        )
        if note_match:
            note = note_match.group(1).strip()
            # Split into sentences and take meaningful ones
            sentences = re.split(r"[.\n]", note)
            for s in sentences:
                s = s.strip()
                if s and len(s) > 5:
                    instructions.append(s)

        # Look for delivery instructions
        delivery_match = re.search(
            r"[Dd]elivery\s*.+?(?:before|by|must|completed).+?(?:\.|$)",
            text,
        )
        if delivery_match:
            instructions.append(delivery_match.group(0).strip())

        return "; ".join(instructions[:3]) if instructions else None

    def _parse_number(self, value: str, as_int: bool = False) -> float | int:
        """Parse number from string."""
        cleaned = re.sub(r"[^\d.\-]", "", str(value))
        try:
            if as_int:
                return int(float(cleaned)) if cleaned else 1
            return float(cleaned) if cleaned else 0.0
        except ValueError:
            return 1 if as_int else 0.0

    def _generate_code(self, description: str) -> str:
        """Generate product code from description."""
        if not description:
            return "ITEM"
        # Extract model number if present
        model_match = re.search(r"Model\s*([A-Z0-9]+)", description, re.IGNORECASE)
        if model_match:
            return model_match.group(1)
        # Otherwise use initials
        words = description.split()[:4]
        return "".join(w[0].upper() for w in words if w and w[0].isalpha()) or "ITEM"
